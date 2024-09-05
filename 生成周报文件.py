import os
import tkinter as tk
from tkinter import messagebox
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# 设置字体为宋体的函数
def set_font_to_simsun(element, font_size=12):
    """设置元素的字体为宋体"""
    run = element.runs[0]
    run.font.name = '宋体'
    run.font.size = Pt(font_size)
    # 设置中文字体支持
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def validate_date(date_str):
    """验证日期格式"""
    try:
        return datetime.strptime(date_str, "%Y-%m-%d")
    except ValueError:
        return None

def create_folder(folder_name):
    """创建文件夹"""
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)

def create_doc(file_path, week_start, week_end, report_content):
    """创建并保存docx文件"""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    heading = doc.add_heading(f"周报：{week_start.strftime('%Y年%m月%d日')} - {week_end.strftime('%Y年%m月%d日')}", level=1)
    set_font_to_simsun(heading, font_size=14)

    paragraph = doc.add_paragraph(report_content)
    set_font_to_simsun(paragraph)
    
    doc.save(file_path)

def generate_reports():
    """生成周报文件"""
    start_date_str = entry_start_date.get()
    end_date_str = entry_end_date.get()
    file_name_prefix = entry_file_name.get().strip()
    report_content = entry_report_content.get("1.0", tk.END).strip()
    folder_name = entry_folder_name.get().strip()
    
    start_date = validate_date(start_date_str)
    end_date = validate_date(end_date_str)
    
    if not start_date or not end_date:
        messagebox.showerror("错误", "请输入正确的日期格式 (YYYY-MM-DD)")
        return
    
    if not file_name_prefix:
        messagebox.showerror("错误", "请输入文件名称前缀")
        return
    
    if not report_content:
        messagebox.showerror("错误", "请输入周报内容")
        return
    
    if not folder_name:
        messagebox.showerror("错误", "请输入文件夹名称")
        return

    create_folder(folder_name)
    
    current_date = start_date
    one_week = timedelta(weeks=1)

    while current_date <= end_date:
        week_start = current_date
        week_end = min(current_date + timedelta(days=4), end_date)
        file_name = f"{file_name_prefix}-{week_start.month}月{week_start.day}日-{week_end.month}月{week_end.day}日.docx"
        file_path = os.path.join(folder_name, file_name)
        
        create_doc(file_path, week_start, week_end, report_content)
        print(f"已生成文件: {file_path}")
        
        current_date += one_week
    
    messagebox.showinfo("完成", "周报文件已生成！")

def create_gui():
    """创建GUI界面"""
    root = tk.Tk()
    root.title("周报生成器")

    tk.Label(root, text="开始日期 (YYYY-MM-DD):").grid(row=0, column=0, padx=10, pady=5, sticky='e')
    tk.Label(root, text="结束日期 (YYYY-MM-DD):").grid(row=1, column=0, padx=10, pady=5, sticky='e')
    tk.Label(root, text="文件名称前缀:").grid(row=2, column=0, padx=10, pady=5, sticky='e')
    tk.Label(root, text="文件夹名称:").grid(row=3, column=0, padx=10, pady=5, sticky='e')
    tk.Label(root, text="周报内容:").grid(row=4, column=0, padx=10, pady=5, sticky='ne')

    global entry_start_date, entry_end_date, entry_file_name, entry_folder_name, entry_report_content

    entry_start_date = tk.Entry(root)
    entry_start_date.grid(row=0, column=1, padx=10, pady=5)

    entry_end_date = tk.Entry(root)
    entry_end_date.grid(row=1, column=1, padx=10, pady=5)

    entry_file_name = tk.Entry(root)
    entry_file_name.grid(row=2, column=1, padx=10, pady=5)

    entry_folder_name = tk.Entry(root)
    entry_folder_name.grid(row=3, column=1, padx=10, pady=5)

    entry_report_content = tk.Text(root, height=5, width=40)
    entry_report_content.grid(row=4, column=1, padx=10, pady=5, sticky='w')

    generate_button = tk.Button(root, text="生成周报", command=generate_reports)
    generate_button.grid(row=5, columnspan=2, padx=10, pady=10)

    root.mainloop()

if __name__ == "__main__":
    create_gui()
